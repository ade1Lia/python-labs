
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document()

doc.add_paragraph('В микроконтроллерах ATmega, которые применяются в Arduino, можно выделить три типа памяти: ')

memories = [
    ("Флеш-память: используется для хранения скетчей", False),
    ("ОЗУ (SRAM — статическая оперативная память с произвольным доступом): используется для работы с переменными.", True),
    ("EEPROM (энергонезависимая память): используется для хранения постоянных данных.", False)
]

for text, is_bold in memories:
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(text).bold = is_bold
    para.paragraph_format.left_indent = Inches(1)

doc.add_paragraph('Флеш-память и EEPROM не теряют данные при отключении питания, в то время как ОЗУ является энергозависимой.')

headers = ['ATmega168', 'ATmega328', 'ATmega1280', 'ATmega2560']
rows = [
    ('Flash (1 кБ занят загрузчиком)', '16 Кбайт', '32 Кбайт', '128 Кбайт', '256 Кбайт'),
    ('SRAM', '1 Кбайт', '2 Кбайт', '8 Кбайт', '8 Кбайт'),
    ('EEPROM', '512 байт', '1024 байта', '4 Кбайта', '4 Кбайта')
]

table = doc.add_table(rows=1, cols=5)

hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    cell = hdr_cells[i + 1]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for row in rows:
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = value
        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            row_cells[i].paragraphs[0].runs[0].font.bold = True

def apply_gray_background(cell):
    cell_xml = cell._tc
    cell_properties = cell_xml.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'eeeeee')
    cell_properties.append(shading)

for row in table.rows[0].cells:
    apply_gray_background(row)
for row in table.rows[1:]:
    apply_gray_background(row.cells[0])

final_paragraph = doc.add_paragraph()
final_paragraph.add_run('По утверждениям производителя, EEPROM имеет гарантию на 100 000 циклов записи/стирания и 100 лет хранения данных при температуре 25°C. Эти данные не относятся к операциям чтения, которые не имеют ограничений. Исходя из этого, рекомендуется осторожно обращаться с EEPROM в своих скетчах.').italic = True
final_paragraph.paragraph_format.space_before = Pt(12)

doc.save('document.docx')
