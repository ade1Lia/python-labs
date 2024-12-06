from docx import Document
from docx.shared import Inches
doc = Document('document.docx')
img_paragraph = doc.add_paragraph()
run = img_paragraph.add_run()
run.add_picture('cat.png', width=Inches(1.25))
caption = doc.add_paragraph('Кисик')
doc.save('document.docx')